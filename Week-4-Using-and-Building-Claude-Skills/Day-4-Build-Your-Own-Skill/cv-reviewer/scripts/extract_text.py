#!/usr/bin/env python3
"""
Text extraction utility for the CV Reviewer skill.

Pulls raw text out of a PDF, DOCX, or plain-text resume so the rest of the
pipeline (helpers.py, review_cv.py) has a single normalized string to work
with, regardless of the source file format.

Usage:
    python3 extract_text.py <path/to/resume.(pdf|docx|txt)>

Optional dependencies (only imported if the matching file type is used):
    pip install pdfplumber python-docx

If a dependency isn't installed, extraction falls back to a clear error
message rather than a stack trace, since the caller (Claude, following
SKILL.md) needs to distinguish "extraction failed for a fixable reason"
from "this document has no text layer at all."
"""

import sys
import json
import os

MIN_TEXT_DENSITY_CHARS_PER_PAGE = 200
# Below this, a PDF is very likely a scanned image with no real text layer -
# the single most common hard ATS failure. See references/ats-best-practices.md.


def extract_from_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def extract_from_pdf(path: str) -> "tuple[str, int]":
    """Returns (text, page_count). Raises RuntimeError with a clear message on failure."""
    try:
        import pdfplumber
    except ImportError:
        try:
            from pypdf import PdfReader
        except ImportError:
            raise RuntimeError(
                "No PDF extraction library available. Install one of: "
                "'pip install pdfplumber' (preferred) or 'pip install pypdf'."
            )
        reader = PdfReader(path)
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages), len(pages)

    with pdfplumber.open(path) as pdf:
        pages = [page.extract_text() or "" for page in pdf.pages]
        return "\n".join(pages), len(pdf.pages)


def extract_from_docx(path: str) -> str:
    try:
        import docx
    except ImportError:
        raise RuntimeError(
            "python-docx is required to read .docx files. Install with: "
            "'pip install python-docx'."
        )
    document = docx.Document(path)
    paragraphs = [p.text for p in document.paragraphs]

    # DOCX resumes frequently put dates/skills in tables - don't silently drop them.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    return "\n".join(paragraphs)


def extract(path: str) -> dict:
    """
    Main entry point. Returns a dict:
        {
            "text": str,
            "format": "pdf" | "docx" | "txt",
            "warnings": [str, ...]
        }
    """
    if not os.path.isfile(path):
        raise FileNotFoundError(f"No such file: {path}")

    ext = os.path.splitext(path)[1].lower()
    warnings = []

    if ext == ".pdf":
        text, page_count = extract_from_pdf(path)
        density = len(text.strip()) / max(page_count, 1)
        if density < MIN_TEXT_DENSITY_CHARS_PER_PAGE:
            warnings.append(
                "LOW TEXT DENSITY DETECTED: this PDF likely has no real text layer "
                "(commonly a scanned image or flattened graphic). This is a critical "
                "ATS Compatibility failure - flag it directly rather than reviewing "
                "the near-empty extracted text as if it were the actual resume content. "
                "See references/ats-best-practices.md."
            )
        fmt = "pdf"

    elif ext == ".docx":
        text = extract_from_docx(path)
        fmt = "docx"

    elif ext in (".txt", ".md"):
        text = extract_from_txt(path)
        fmt = "txt"

    else:
        raise ValueError(
            f"Unsupported file extension '{ext}'. Expected .pdf, .docx, or .txt."
        )

    if not text.strip():
        warnings.append(
            "Extracted text is empty. The document may be corrupted, password-protected, "
            "or entirely image-based. Ask the user to paste the resume text directly."
        )

    return {"text": text, "format": fmt, "warnings": warnings}


def main():
    if len(sys.argv) != 2:
        print("Usage: python3 extract_text.py <path/to/resume.(pdf|docx|txt)>", file=sys.stderr)
        sys.exit(1)

    path = sys.argv[1]
    try:
        result = extract(path)
    except (FileNotFoundError, ValueError, RuntimeError) as e:
        print(json.dumps({"error": str(e)}), file=sys.stderr)
        sys.exit(1)

    for warning in result["warnings"]:
        print(f"WARNING: {warning}", file=sys.stderr)

    print(result["text"])


if __name__ == "__main__":
    main()
